"""生成 Word / PDF / Excel 穷举样例（文/表/图组合）。"""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
SAMPLES = ROOT / "samples"
ASSETS = ROOT / "assets"


def _font(size: int = 20):
    for name in (
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ):
        p = Path(name)
        if p.exists():
            try:
                return ImageFont.truetype(str(p), size)
            except OSError:
                pass
    return ImageFont.load_default()


def make_assets() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = {}
    specs = [
        ("img_north.png", "NORTH_IMG", (37, 99, 235), (220, 80)),
        ("img_south.png", "SOUTH_IMG", (5, 150, 105), (220, 80)),
        ("img_chart.png", "CHART_IMG", (220, 38, 38), (280, 120)),
        ("img_logo.png", "LOGO_IMG", (100, 100, 100), (120, 60)),
    ]
    font = _font(22)
    for fname, label, color, size in specs:
        path = ASSETS / fname
        im = Image.new("RGB", size, color)
        dr = ImageDraw.Draw(im)
        dr.text((12, size[1] // 2 - 12), label, fill=(255, 255, 255), font=font)
        im.save(path)
        out[label] = path
    return out


def _add_para(doc, text, *, bold=False, style=None):
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
    return p


def gen_docx(assets: dict[str, Path]) -> list[Path]:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    SAMPLES.mkdir(parents=True, exist_ok=True)
    paths = []

    def save(name: str, build) -> Path:
        p = SAMPLES / name
        doc = Document()
        build(doc)
        doc.save(p)
        paths.append(p)
        return p

    # 1 纯文字
    def t_text(doc):
        doc.add_heading("纯文字报告 WORD_TEXT_ONLY", level=1)
        doc.add_heading("概述", level=2)
        _add_para(doc, "华北增长 12%。海外下滑 5%。WORD_BODY_MARKER")
        _add_para(doc, "重点客户跟进", style="List Bullet")
        _add_para(doc, "渠道漏斗复盘", style="List Bullet")
        _add_para(doc, "输出行动清单", style="List Number")
        p = doc.add_paragraph()
        r = p.add_run("红色强调：WORD_RED_MARKER")
        r.bold = True
        r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    save("W01_text_only.docx", t_text)

    # 2 文字+单表
    def t_text_table(doc):
        doc.add_heading("文字加单表 WORD_TEXT_TABLE", level=1)
        _add_para(doc, "下表为区域收入。WORD_TT_MARKER")
        t = doc.add_table(rows=4, cols=3)
        t.style = "Table Grid"
        data = [("区域", "营收", "同比"), ("华北", "1280", "+12%"), ("华南", "960", "0%"), ("海外", "430", "-5%")]
        for i, row in enumerate(data):
            for j, v in enumerate(row):
                t.rows[i].cells[j].text = v

    save("W02_text_table.docx", t_text_table)

    # 3 文字+单图
    def t_text_img(doc):
        doc.add_heading("文字加单图 WORD_TEXT_IMG", level=1)
        _add_para(doc, "见下图示意。WORD_TI_MARKER")
        doc.add_picture(str(assets["NORTH_IMG"]), width=Inches(2.2))
        _add_para(doc, "图注：华北示意图 CAPTION_NORTH")

    save("W03_text_image.docx", t_text_img)

    # 4 文字+图+表
    def t_all(doc):
        doc.add_heading("综合：文图表 WORD_MIXED", level=1)
        _add_para(doc, "综合样例正文 WORD_MIX_MARKER")
        doc.add_picture(str(assets["CHART_IMG"]), width=Inches(2.5))
        _add_para(doc, "图注：CHART_CAPTION")
        t = doc.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        for i, row in enumerate([("指标", "值"), ("准确率", "96%"), ("召回率", "91%")]):
            for j, v in enumerate(row):
                t.rows[i].cells[j].text = v
        _add_para(doc, "结尾段 WORD_MIX_END")

    save("W04_text_image_table.docx", t_all)

    # 5 多表
    def t_multi_table(doc):
        doc.add_heading("多表 WORD_MULTI_TABLE", level=1)
        _add_para(doc, "表一收入 WORD_MT1")
        t1 = doc.add_table(rows=3, cols=2)
        t1.style = "Table Grid"
        for i, row in enumerate([("区", "值"), ("华北", "1280"), ("海外", "430")]):
            for j, v in enumerate(row):
                t1.rows[i].cells[j].text = v
        _add_para(doc, "表二成本 WORD_MT2")
        t2 = doc.add_table(rows=3, cols=2)
        t2.style = "Table Grid"
        for i, row in enumerate([("项", "额"), ("人力", "200"), ("云资源", "80")]):
            for j, v in enumerate(row):
                t2.rows[i].cells[j].text = v

    save("W05_multi_table.docx", t_multi_table)

    # 6 多图
    def t_multi_img(doc):
        doc.add_heading("多图 WORD_MULTI_IMG", level=1)
        _add_para(doc, "多图样例 WORD_MI_MARKER")
        doc.add_picture(str(assets["NORTH_IMG"]), width=Inches(1.8))
        _add_para(doc, "图1 CAPTION_N1")
        doc.add_picture(str(assets["SOUTH_IMG"]), width=Inches(1.8))
        _add_para(doc, "图2 CAPTION_S1")
        doc.add_picture(str(assets["LOGO_IMG"]), width=Inches(1.2))
        _add_para(doc, "图3 CAPTION_L1")

    save("W06_multi_image.docx", t_multi_img)

    # 7 合并单元格表
    def t_merge(doc):
        doc.add_heading("合并单元格 WORD_MERGE_TABLE", level=1)
        t = doc.add_table(rows=3, cols=3)
        t.style = "Table Grid"
        t.cell(0, 0).text = "合计区 MERGE_HDR"
        t.cell(0, 0).merge(t.cell(0, 2))
        t.cell(1, 0).text = "华北"
        t.cell(1, 1).text = "1280"
        t.cell(1, 2).text = "+12%"
        t.cell(2, 0).text = "海外"
        t.cell(2, 1).text = "430"
        t.cell(2, 2).text = "-5%"

    save("W07_merged_table.docx", t_merge)

    # 8 页眉页脚+正文
    def t_hf(doc):
        sec = doc.sections[0]
        sec.header.paragraphs[0].text = "页眉 HEADER_WORD"
        sec.footer.paragraphs[0].text = "页脚 FOOTER_WORD"
        doc.add_heading("含页眉页脚 WORD_HF", level=1)
        _add_para(doc, "正文 BODY_WORD_HF")

    save("W08_header_footer.docx", t_hf)

    # 9 超链接
    def t_link(doc):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc.add_heading("超链接 WORD_LINK", level=1)
        p = doc.add_paragraph()
        part = p.part
        r_id = part.relate_to(
            "https://example.com/office",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = "外部链接 LINK_WORD"
        new_run.append(text_el)
        hyperlink.append(new_run)
        p._p.append(hyperlink)
        _add_para(doc, "链接后正文 WORD_LINK_BODY")

    save("W09_hyperlink.docx", t_link)

    # 10 多级标题长文
    def t_long(doc):
        doc.add_heading("长文结构 WORD_LONG", level=1)
        for i in range(1, 4):
            doc.add_heading(f"第{i}章 CHAPTER_{i}", level=2)
            doc.add_heading(f"小节 {i}.1 SEC_{i}", level=3)
            _add_para(doc, f"章节正文段落 {i} WORD_LONG_P{i}")

    save("W10_long_headings.docx", t_long)

    return paths


def gen_xlsx(assets: dict[str, Path]) -> list[Path]:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font, PatternFill

    paths = []

    def save(name, build):
        p = SAMPLES / name
        wb = Workbook()
        build(wb)
        wb.save(p)
        paths.append(p)
        return p

    def t_text(wb):
        ws = wb.active
        ws.title = "说明"
        ws["A1"] = "纯文字工作簿 EXCEL_TEXT_ONLY"
        ws["A2"] = "华北增长 12%。EXCEL_BODY_MARKER"
        ws["A3"] = "行动：跟进重点客户"

    save("E01_text_only.xlsx", t_text)

    def t_text_table(wb):
        ws = wb.active
        ws.title = "收入"
        ws["A1"] = "文字加表 EXCEL_TEXT_TABLE"
        ws["A2"] = "说明 EXCEL_TT_MARKER"
        for c, h in enumerate(["区域", "营收", "同比"], 1):
            cell = ws.cell(4, c, h)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="2563EB")
        for i, row in enumerate([("华北", 1280, "+12%"), ("华南", 960, "0%"), ("海外", 430, "-5%")], 5):
            for j, v in enumerate(row, 1):
                ws.cell(i, j, v)

    save("E02_text_table.xlsx", t_text_table)

    def t_text_img(wb):
        ws = wb.active
        ws.title = "图示"
        ws["A1"] = "文字加图 EXCEL_TEXT_IMG"
        ws["A2"] = "EXCEL_TI_MARKER"
        try:
            img = XLImage(str(assets["NORTH_IMG"]))
            img.width = 160
            img.height = 60
            ws.add_image(img, "A4")
            ws["A8"] = "图下方说明 CAPTION_EXCEL_NORTH"
        except Exception:
            ws["A4"] = "（图片嵌入失败，仅文本） CAPTION_EXCEL_NORTH"

    save("E03_text_image.xlsx", t_text_img)

    def t_mix(wb):
        ws = wb.active
        ws.title = "综合"
        ws["A1"] = "文图表 EXCEL_MIXED"
        ws["A2"] = "EXCEL_MIX_MARKER"
        try:
            img = XLImage(str(assets["CHART_IMG"]))
            img.width = 180
            img.height = 80
            ws.add_image(img, "A3")
        except Exception:
            pass
        ws["A10"] = "指标"
        ws["B10"] = "值"
        ws["A11"] = "准确率"
        ws["B11"] = "96%"
        ws["A12"] = "召回率"
        ws["B12"] = "91%"

    save("E04_text_image_table.xlsx", t_mix)

    def t_multi_table(wb):
        ws = wb.active
        ws.title = "多表"
        ws["A1"] = "左表 EXCEL_MT_LEFT"
        ws["A2"] = "区"
        ws["B2"] = "值"
        ws["A3"] = "华北"
        ws["B3"] = 1280
        ws["D1"] = "右表 EXCEL_MT_RIGHT"
        ws["D2"] = "项"
        ws["E2"] = "额"
        ws["D3"] = "人力"
        ws["E3"] = 200
        # 中间留空列 C 作为分隔

    save("E05_multi_table_side.xlsx", t_multi_table)

    def t_multi_sheet(wb):
        ws1 = wb.active
        ws1.title = "华北"
        ws1["A1"] = "SHEET_NORTH"
        ws1["A2"] = "营收"
        ws1["B2"] = 1280
        ws2 = wb.create_sheet("海外")
        ws2["A1"] = "SHEET_OVERSEAS"
        ws2["A2"] = "营收"
        ws2["B2"] = 430
        ws3 = wb.create_sheet("汇总")
        ws3["A1"] = "SHEET_SUMMARY"
        ws3["A2"] = "合计"
        ws3["B2"] = 1710

    save("E06_multi_sheet.xlsx", t_multi_sheet)

    def t_merged(wb):
        ws = wb.active
        ws.title = "合并"
        ws.merge_cells("A1:C1")
        ws["A1"] = "合并表头 EXCEL_MERGE_HDR"
        ws["A2"] = "华北"
        ws["B2"] = 1280
        ws["C2"] = "+12%"
        ws["A3"] = "海外"
        ws["B3"] = 430
        ws["C3"] = "-5%"

    save("E07_merged.xlsx", t_merged)

    def t_formats(wb):
        ws = wb.active
        ws.title = "格式"
        ws["A1"] = "EXCEL_FMT_MARKER"
        ws["A2"] = "百分比"
        ws["B2"] = 0.12
        ws["B2"].number_format = "0%"
        ws["A3"] = "金额"
        ws["B3"] = 1280
        ws["B3"].number_format = '"¥"#,##0'
        ws["A4"] = "红字"
        ws["B4"] = -5
        ws["B4"].font = Font(color="DC2626", bold=True)

    save("E08_number_formats.xlsx", t_formats)

    def t_formula(wb):
        ws = wb.active
        ws.title = "公式"
        ws["A1"] = "EXCEL_FORMULA_MARKER"
        ws["A2"] = 100
        ws["B2"] = 200
        ws["C2"] = "=A2+B2"

    save("E09_formula.xlsx", t_formula)

    def t_multi_img(wb):
        ws = wb.active
        ws.title = "多图"
        ws["A1"] = "EXCEL_MULTI_IMG"
        try:
            for anchor, key in [("A3", "NORTH_IMG"), ("A8", "SOUTH_IMG")]:
                img = XLImage(str(assets[key]))
                img.width = 140
                img.height = 50
                ws.add_image(img, anchor)
            ws["A12"] = "双图说明 EXCEL_MI_CAPTION"
        except Exception:
            ws["A3"] = "EXCEL_MI_CAPTION"

    save("E10_multi_image.xlsx", t_multi_img)

    return paths


def gen_pdf(assets: dict[str, Path]) -> list[Path]:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image as RLImage,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    paths = []
    styles = getSampleStyleSheet()

    def save(name, flow):
        p = SAMPLES / name
        doc = SimpleDocTemplate(str(p), pagesize=A4)
        doc.build(flow())
        paths.append(p)
        return p

    def t_text():
        return [
            Paragraph("PDF Text Only PDF_TEXT_ONLY", styles["Title"]),
            Spacer(1, 12),
            Paragraph("North China growth 12%. Overseas -5%. PDF_BODY_MARKER", styles["Normal"]),
            Paragraph("1. Follow key accounts", styles["Normal"]),
            Paragraph("2. Review funnel PDF_LIST_MARKER", styles["Normal"]),
        ]

    save("P01_text_only.pdf", t_text)

    def t_text_table():
        data = [
            ["Region", "Revenue", "YoY"],
            ["North", "1280", "+12%"],
            ["South", "960", "0%"],
            ["Overseas", "430", "-5%"],
        ]
        t = Table(data)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]
            )
        )
        return [
            Paragraph("PDF Text+Table PDF_TEXT_TABLE", styles["Title"]),
            Paragraph("Intro PDF_TT_MARKER", styles["Normal"]),
            Spacer(1, 10),
            t,
        ]

    save("P02_text_table.pdf", t_text_table)

    def t_text_img():
        return [
            Paragraph("PDF Text+Image PDF_TEXT_IMG", styles["Title"]),
            Paragraph("See figure PDF_TI_MARKER", styles["Normal"]),
            Spacer(1, 8),
            RLImage(str(assets["NORTH_IMG"]), width=2.2 * inch, height=0.8 * inch),
            Paragraph("Caption CAPTION_PDF_NORTH", styles["Normal"]),
        ]

    save("P03_text_image.pdf", t_text_img)

    def t_mix():
        data = [["Metric", "Value"], ["Accuracy", "96%"], ["Recall", "91%"]]
        t = Table(data)
        t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]))
        return [
            Paragraph("PDF Mixed PDF_MIXED", styles["Title"]),
            Paragraph("Body PDF_MIX_MARKER", styles["Normal"]),
            RLImage(str(assets["CHART_IMG"]), width=2.5 * inch, height=1.0 * inch),
            Paragraph("Caption CHART_PDF_CAPTION", styles["Normal"]),
            Spacer(1, 8),
            t,
            Paragraph("End PDF_MIX_END", styles["Normal"]),
        ]

    save("P04_text_image_table.pdf", t_mix)

    def t_multi_table():
        t1 = Table([["A", "B"], ["North", "1280"], ["Overseas", "430"]])
        t1.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]))
        t2 = Table([["Item", "Cost"], ["Labor", "200"], ["Cloud", "80"]])
        t2.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]))
        return [
            Paragraph("PDF Multi Tables PDF_MULTI_TABLE", styles["Title"]),
            Paragraph("Table1 PDF_MT1", styles["Normal"]),
            t1,
            Spacer(1, 12),
            Paragraph("Table2 PDF_MT2", styles["Normal"]),
            t2,
        ]

    save("P05_multi_table.pdf", t_multi_table)

    def t_multi_img():
        return [
            Paragraph("PDF Multi Images PDF_MULTI_IMG", styles["Title"]),
            Paragraph("PDF_MI_MARKER", styles["Normal"]),
            RLImage(str(assets["NORTH_IMG"]), width=1.8 * inch, height=0.65 * inch),
            Paragraph("Cap1 CAPTION_PDF_N1", styles["Normal"]),
            RLImage(str(assets["SOUTH_IMG"]), width=1.8 * inch, height=0.65 * inch),
            Paragraph("Cap2 CAPTION_PDF_S1", styles["Normal"]),
            RLImage(str(assets["LOGO_IMG"]), width=1.2 * inch, height=0.55 * inch),
            Paragraph("Cap3 CAPTION_PDF_L1", styles["Normal"]),
        ]

    save("P06_multi_image.pdf", t_multi_img)

    def t_multipage():
        story = [Paragraph("PDF Multipage PDF_MULTIPAGE", styles["Title"])]
        for i in range(1, 4):
            story += [
                Paragraph(f"Section {i} PDF_PAGE_{i}", styles["Heading1"]),
                Paragraph(f"Content block {i}.", styles["Normal"]),
                PageBreak(),
            ]
        story.append(Paragraph("Last page PDF_PAGE_LAST", styles["Normal"]))
        return story

    save("P07_multipage.pdf", t_multipage)

    def t_scanned():
        # 整页图片模拟扫描件
        from reportlab.pdfgen import canvas

        p = SAMPLES / "P08_scanned_like.pdf"
        img = Image.new("RGB", (800, 1100), (255, 255, 255))
        dr = ImageDraw.Draw(img)
        dr.text((60, 200), "SCANNED_PDF_MARKER", fill=(0, 0, 0), font=_font(28))
        dr.text((60, 280), "Image-only page (no text layer)", fill=(0, 0, 0), font=_font(18))
        tmp = ASSETS / "_scan_page.png"
        img.save(tmp)
        c = canvas.Canvas(str(p), pagesize=A4)
        c.drawImage(str(tmp), 40, 80, width=500, height=680)
        c.save()
        paths.append(p)
        return p

    t_scanned()

    def t_two_col_like():
        # 用双列表格模拟双栏
        left = Paragraph("Left column PDF_LEFT_COL North revenue discussion.", styles["Normal"])
        right = Paragraph("Right column PDF_RIGHT_COL Overseas review notes.", styles["Normal"])
        t = Table([[left, right]], colWidths=[3.2 * inch, 3.2 * inch])
        t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
        return [
            Paragraph("PDF Two Column PDF_TWO_COL", styles["Title"]),
            t,
        ]

    save("P09_two_column.pdf", t_two_col_like)

    return paths


def main():
    assets = make_assets()
    w = gen_docx(assets)
    e = gen_xlsx(assets)
    p = gen_pdf(assets)
    print(f"generated word={len(w)} excel={len(e)} pdf={len(p)} total={len(w)+len(e)+len(p)}")


if __name__ == "__main__":
    main()
