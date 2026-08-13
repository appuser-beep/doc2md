"""文档转 Markdown 转换封装。"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO, Callable

from cleanup import clean_markdown, clean_markdown_light

APP_VERSION = "1.7.9"

_BUILTIN_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".pptx",
    ".xlsx",
    ".xls",
    ".msg",
    ".epub",
    ".ipynb",
    ".jpg",
    ".jpeg",
    ".png",
    ".wav",
    ".mp3",
    ".m4a",
    ".mp4",
    ".html",
    ".htm",
    ".csv",
    ".json",
    ".jsonl",
    ".xml",
    ".rss",
    ".atom",
    ".txt",
    ".text",
    ".md",
    ".markdown",
    ".zip",
}

_CU_EXTENSIONS = {
    ".eml",
    ".rtf",
    ".heif",
    ".heic",
    ".bmp",
    ".tif",
    ".tiff",
    ".flac",
    ".ogg",
    ".aac",
    ".wma",
    ".m4v",
    ".mov",
    ".avi",
    ".mkv",
    ".webm",
    ".flv",
    ".wmv",
}

SUPPORTED_EXTENSIONS = _BUILTIN_EXTENSIONS | _CU_EXTENSIONS

UNSUPPORTED_BUT_RISKY = {
    ".doc": "老版 Word（.doc）不受支持，请先另存为 .docx。若强行转换，可能只得到乱码片段。",
    ".ppt": "老版 PowerPoint（.ppt）不受支持，请先另存为 .pptx。",
    ".xlsb": "不支持 .xlsb，请另存为 .xlsx。",
}

FILE_DIALOG_TYPES = [
    (
        "支持的文档",
        "*.pdf *.docx *.pptx *.xlsx *.xls *.msg *.eml *.epub *.ipynb "
        "*.jpg *.jpeg *.png *.bmp *.tif *.tiff *.heif *.heic "
        "*.wav *.mp3 *.m4a *.mp4 *.flac *.ogg *.aac *.wma "
        "*.mov *.avi *.mkv *.webm *.m4v *.flv *.wmv "
        "*.html *.htm *.rtf *.csv *.json *.jsonl *.xml *.rss *.atom *.txt *.md *.zip",
    ),
    ("PDF", "*.pdf"),
    ("Word", "*.docx"),
    ("PowerPoint", "*.pptx"),
    ("Excel", "*.xlsx *.xls"),
    ("Outlook / 邮件", "*.msg *.eml"),
    ("图片", "*.jpg *.jpeg *.png *.bmp *.tif *.tiff *.heif *.heic"),
    ("音视频", "*.wav *.mp3 *.m4a *.mp4 *.flac *.ogg *.aac *.wma *.mov *.avi *.mkv *.webm"),
    ("网页/文本", "*.html *.htm *.txt *.md *.rtf *.csv *.json *.xml"),
    ("压缩包", "*.zip"),
    ("所有文件", "*.*"),
]

HELP_TAB_FORMATS = """支持格式与能力
版本 {version}

【本地即可用】
• Office：Word (.docx)、Excel (.xlsx / .xls)、PowerPoint (.pptx)、Outlook (.msg)
• 文档：PDF、EPUB、Jupyter Notebook (.ipynb)
• 数据：HTML、RSS、Atom、CSV、JSON、JSONL、XML、TXT、Markdown
• 图片：JPG / PNG（可读 EXIF 元数据）
• 音频：WAV / MP3 / M4A（语音转写，通常需联网）
• 视频：MP4 音轨转写（建议安装 ffmpeg）
• 压缩包：标准 ZIP（自动遍历包内文档与源码）
• 网络：网页、Wikipedia、Bing 搜索、YouTube 字幕

【需 Azure 云端】（在「Azure 设置」配置）
• 文档智能：扫描件 PDF、复杂版面、BMP / TIFF 等
• 内容理解：视频、EML、RTF、更多音频格式

【需大模型】（在「大模型设置」配置）
• 图片内容描述、PPT 内嵌图说明
• OCR 插件：识别 PDF / Office 内嵌图中的文字

【默认本地增强】（未勾选「窄接口模式」时）
• Excel：合并单元格、宽表整理、日期/百分比/货币可读化
• Notebook：保留代码输出（含 Markdown / 图片占位）
• ZIP：校验真实格式、跳过二进制与压缩包内嵌包
• Word：补充页眉、页脚文字
""".format(version=APP_VERSION)

HELP_TAB_GLOSSARY = """名词解释（设置里常见英文）

API Key / 密钥
  调用云服务或大模型的通行证，相当于密码。请勿发给他人或提交到公开仓库。

Base URL / 接口地址
  大模型服务的访问网址。留空通常走默认网关；使用代理或私有部署时再填写。

llm_model / 模型名称
  具体用哪一个大模型，例如 gpt-4o。名称需与服务商控制台一致。

llm_prompt / 提示词
  告诉模型「怎么描述图片」的说明文字。可按业务改成更具体的要求。

OCR（光学字符识别）
  把图片或扫描件里的文字识别成可编辑文本。纯文字 PDF 一般不需要。

Endpoint（终结点 / 接入地址）
  Azure 服务的专属网址，形如 https://xxxx.cognitiveservices.azure.com/

Document Intelligence / 文档智能
  Azure 能力：抽取扫描 PDF、复杂版面、BMP/TIFF 等中的文字与结构。

docintel_endpoint / docintel_api_version / docintel_file_types
  文档智能的接入地址、接口版本、处理的文件扩展名（逗号分隔，如 pdf,docx）。
  留空则使用软件内置默认类型。

Content Understanding / 内容理解
  Azure 能力：处理视频、邮件 EML、RTF、扩展音频等。

cu_endpoint / cu_analyzer_id / cu_file_types
  内容理解的接入地址、分析器编号、文件类型列表。

ExifTool
  第三方小工具，用于读取图片/音频的拍摄时间、设备等元数据。留空会自动查找。

样式映射（style map）
  把 Word 段落样式对应到 Markdown 标题等级，例如：
  p[style-name='Heading 1'] => h1

窄接口模式
  关闭 Excel / Notebook / ZIP 本地增强，更接近最基础的转换路径，便于对照排查。

保留内嵌图 / keep_data_uris / base64
  默认会缩短超长图片数据以减小 Markdown 体积；勾选后保留完整图片（文件变大）。

自定义插件
  自行编写的 .py 脚本，需提供 register_converters(...) 以扩展转换能力。
"""

HELP_TAB_SETTINGS = """设置怎么配（建议顺序）

1）日常办公文档（Word / Excel / PPT / 文本 PDF）
  • 无需改设置，直接转换即可。
  • 颜色、字号不会进入 Markdown，这是格式本身的限制。

2）扫描件 PDF、图片里的字看不清
  • 优先：大模型设置 → 启用 OCR 插件
  • 或：Azure 设置 → 启用文档智能，并填写 Endpoint / Key

3）需要给图片写说明（无障碍读图）
  • 大模型设置 → 启用大模型，填写 API 密钥与模型名称

4）视频、EML、RTF、特殊音频
  • Azure 设置 → 启用内容理解，填写 cu_endpoint 等

5）高级用户
  • 高级设置：样式映射、ExifTool、自定义插件、窄接口、保留内嵌图
  • 命令行：doc2md-cli --list-plugins 查看已加载插件

配置文件位置（一般无需手改）
  用户目录下的 .doc2md 文件夹。
"""

HELP_TAB_FAQ = """常见问题

Q：转换后颜色、字号没了？
A：Markdown 只保留结构（标题、列表、表格、链接），不保留视觉排版。

Q：扫描件 PDF 几乎没文字？
A：没有文字层时本地无法“看图识字”。请启用 OCR 或 Azure 文档智能。

Q：Excel 公式变成了 =A1+B1？
A：文件里没有保存计算结果时，只能保留公式文本。请先在 Excel 中打开并保存后再转。

Q：ZIP 提示实际是 RAR？
A：扩展名是 .zip，内容可能是 RAR/7z。请用 7-Zip 重新打成真正的 ZIP。

Q：输出里图片变成占位符？
A：默认不落盘二进制图片。勾选「保留内嵌图」可保留 base64；OCR/大模型可补充文字说明。

Q：图形版和命令行有何区别？
A：转换能力相同。图形版方便点选；命令行适合批量与脚本。

不支持
  老格式 .doc / .ppt / .xlsb；直接转 RAR/7z；Markdown 反向导出为 Office。
"""

SUPPORT_HELP = "\n\n".join(
    [
        HELP_TAB_FORMATS.strip(),
        "──────── 名词解释 ────────",
        HELP_TAB_GLOSSARY.strip(),
        "──────── 设置指南 ────────",
        HELP_TAB_SETTINGS.strip(),
        "──────── 常见问题 ────────",
        HELP_TAB_FAQ.strip(),
    ]
)


class ConversionError(Exception):
    """转换失败。"""


def _result_text(result) -> str:
    text = getattr(result, "markdown", None) or getattr(result, "text_content", None)
    if text is None:
        raise ConversionError("转换完成，但未得到 Markdown 内容。")
    return text


def _get_markitdown():
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise ConversionError(
            "未安装转换引擎。请执行：pip install -r requirements.txt"
        ) from exc

    kwargs: dict = {"enable_plugins": False}
    try:
        from llm_settings import markitdown_llm_kwargs

        kwargs.update(markitdown_llm_kwargs())
    except Exception:
        pass
    try:
        from azure_settings import markitdown_azure_kwargs

        kwargs.update(markitdown_azure_kwargs())
    except Exception:
        pass
    try:
        from advanced_settings import markitdown_advanced_kwargs

        kwargs.update(markitdown_advanced_kwargs())
    except Exception:
        pass

    md = MarkItDown(**kwargs)
    try:
        from plugin_loader import apply_custom_converter_plugins

        apply_custom_converter_plugins(md, kwargs)
    except RuntimeError:
        raise
    except Exception:
        pass
    return md


def _convert_options() -> dict:
    try:
        from advanced_settings import convert_options

        return convert_options()
    except Exception:
        return {}


def _resolve_keep_data_uris(flag: bool | None) -> bool:
    if flag is True:
        return True
    if flag is False:
        return False
    return bool(_convert_options().get("keep_data_uris"))


def _merge_convert_opts(keep_data_uris: bool | None = None) -> dict:
    opts = dict(_convert_options())
    if keep_data_uris is True:
        opts["keep_data_uris"] = True
    elif keep_data_uris is False:
        opts.pop("keep_data_uris", None)
    return opts


def is_hard_block_warning(warn: str | None) -> bool:
    """伪压缩包、空文件等应直接失败，不可「继续尝试」。"""
    if not warn:
        return False
    keys = (
        "实际是 RAR",
        "实际是 7z",
        "不支持 RAR",
        "不支持 7z",
        "文件为空",
    )
    return any(k in warn for k in keys)


def enforce_precheck(source: str) -> None:
    """对硬阻断类预检直接抛错（供 CLI / 窄接口共用）。"""
    warn = precheck_source(source)
    if warn and is_hard_block_warning(warn):
        raise ConversionError(warn)


def precheck_source(source: str) -> str | None:
    source = (source or "").strip()
    if not source or source.startswith(("http://", "https://")):
        return None
    path = Path(source)
    if path.is_file() and path.stat().st_size == 0:
        return "文件为空（0 字节），无法转换。"
    ext = path.suffix.lower()
    if ext in UNSUPPORTED_BUT_RISKY:
        return UNSUPPORTED_BUT_RISKY[ext]
    if ext == ".rar":
        return "不支持 RAR。请用 7-Zip/WinRAR 重新压缩为 ZIP，或解压后选文件转换。"
    if ext == ".7z":
        return "不支持 7z。请改为 ZIP，或解压后选文件转换。"
    if ext == ".zip" and path.is_file():
        try:
            from zip_convert import sniff_archive_kind

            kind = sniff_archive_kind(path)
        except Exception:
            kind = "unknown"
        if kind == "rar":
            return (
                "扩展名是 .zip，但文件实际是 RAR（文件头 Rar!）。"
                "请重新压缩为真正的 ZIP 后再转。"
            )
        if kind == "7z":
            return "扩展名是 .zip，但文件实际是 7z。请改为真正的 ZIP。"
        if kind not in {"zip", "unknown"}:
            return f"扩展名是 .zip，但检测到格式为 {kind}，可能无法转换。"
    if ext and ext not in SUPPORTED_EXTENSIONS:
        return f"扩展名 {ext} 不在常规支持列表中，将尝试自动识别；结果可能不理想。"
    if path.is_file():
        try:
            from azure_settings import precheck_azure_extension

            azure_warn = precheck_azure_extension(path)
            if azure_warn:
                return azure_warn
        except Exception:
            pass
    return None


def _friendly_office_error(path: Path, exc: Exception) -> str | None:
    msg = str(exc)
    low = msg.lower()
    ext = path.suffix.lower()
    if "password" in low or "encrypt" in low:
        return f"文件可能已加密，无法直接转换（{ext}）。请先解除密码后重试。"
    if "badzipfile" in low or "not a zip" in low or "file is not a zip" in low:
        if ext == ".zip":
            try:
                from zip_convert import sniff_archive_kind

                kind = sniff_archive_kind(path)
            except Exception:
                kind = "unknown"
            if kind == "rar":
                return (
                    "扩展名是 .zip，但文件实际是 RAR（文件头 Rar!）。"
                    "请重新压缩为真正的 ZIP 后再转。"
                )
            if kind == "7z":
                return "扩展名是 .zip，但文件实际是 7z。请改为真正的 ZIP。"
            return "ZIP 文件已损坏或不是有效压缩包。"
        return f"文件已损坏或不是有效的 Office/OpenXML 包（{ext}）。"
    if "notolefile" in low or "not an ole2" in low:
        return "不是有效的 Outlook .msg（OLE）文件。"
    if ext in {".doc", ".ppt"} and ("no converter" in low or "could not convert" in low):
        return f"不支持老格式 {ext}，请另存为新格式后重试。"
    if ext in {".wav", ".mp3", ".m4a", ".mp4"} and (
        "recognition" in low or "speech" in low or "transcri" in low or "connection" in low
    ):
        return "音频转写失败（可能需联网语音识别或本地 Whisper）。"
    if ext in {".bmp", ".tif", ".tiff", ".heif", ".heic"} and (
        "no converter" in low or "could not convert" in low
    ):
        return f"内置不支持 {ext}，请在「Azure 设置」启用 Document Intelligence 或 Content Understanding。"
    if ext in {".eml", ".rtf"} and ("no converter" in low or "could not convert" in low):
        return f"{ext} 需 Azure Content Understanding，请在「Azure 设置」配置 cu_endpoint。"
    if ext in {".avi", ".mkv", ".mov", ".webm", ".flv", ".wmv", ".m4v"} and (
        "no converter" in low or "could not convert" in low
    ):
        return f"{ext} 视频需 Azure Content Understanding（cu_endpoint）。"
    return None


def _meaningful_text_len(text: str) -> int:
    """去掉常见 Markdown 图片/空白后，估算可读正文长度。"""
    s = text or ""
    s = re.sub(r"!\[.*?\]\([^)]*\)", "", s)
    s = re.sub(r"<!--.*?-->", "", s, flags=re.DOTALL)
    s = re.sub(r"\s+", "", s)
    return len(s)


def postcheck_result(source: str, text: str) -> str | None:
    stripped = (text or "").strip()
    lower = source.lower()
    meaningful = _meaningful_text_len(stripped)
    if not stripped or meaningful == 0:
        if any(lower.endswith(x) for x in (".jpg", ".jpeg", ".png")):
            return (
                "转换完成，但几乎没有文本。\n"
                "图片默认不做 OCR；可在「大模型设置」启用视觉模型，或启用 OCR 插件。"
            )
        if lower.endswith(".pdf"):
            return (
                "转换完成，但几乎没有文本。\n"
                "若是扫描件/纯图片 PDF，需大模型 OCR 或 Azure Document Intelligence。"
            )
        return "转换完成，但结果为空。"
    if meaningful < 8 and any(lower.endswith(x) for x in (".jpg", ".jpeg", ".png")):
        return "图片输出很少，通常仅有元数据；正文 OCR 需额外能力。"
    # 扫描件常只抽出极少噪声/页眉；阈值过短时同样提示 OCR
    if lower.endswith(".pdf") and meaningful < 12:
        return (
            "转换结果几乎没有可读正文。\n"
            "若是扫描件/纯图片 PDF，请启用 OCR 插件、大模型 OCR 或 Azure Document Intelligence。"
        )
    return None


def _enrich_docx_chrome(path: Path, text: str) -> str:
    """补充 Word 页眉页脚文字。"""
    try:
        from docx import Document
    except Exception:
        return text
    try:
        doc = Document(str(path))
    except Exception:
        return text

    headers: list[str] = []
    footers: list[str] = []
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            t = (p.text or "").strip()
            if t and t not in headers and t not in text:
                headers.append(t)
        for p in sec.footer.paragraphs:
            t = (p.text or "").strip()
            if t and t not in footers and t not in text:
                footers.append(t)
    if not headers and not footers:
        return text

    parts = [text.rstrip()]
    if headers:
        parts.append("\n## 页眉\n\n" + "\n\n".join(headers))
    if footers:
        parts.append("\n## 页脚\n\n" + "\n\n".join(footers))
    return "\n".join(parts).rstrip() + "\n"


def _finalize_markdown(path: Path | None, text: str, *, light: bool = True) -> str:
    cleaner = clean_markdown_light if light else clean_markdown
    text = cleaner(text)
    if light and path is not None and path.suffix.lower() == ".docx":
        text = _enrich_docx_chrome(path, text)
        text = clean_markdown_light(text)
    return text


def _should_use_local_only(local_only: bool | None) -> bool:
    if local_only is not None:
        return local_only
    try:
        from advanced_settings import load_settings

        return bool(load_settings().use_convert_local)
    except Exception:
        return False


def _run_markitdown_local(
    path: Path,
    *,
    progress: Callable[[str], None] | None = None,
    keep_data_uris: bool | None = None,
) -> str:
    if progress:
        progress("正在加载转换引擎…")
    md = _get_markitdown()
    opts = _merge_convert_opts(keep_data_uris)
    try:
        if progress:
            progress(f"正在转换（窄接口）：{path.name}")
        result = md.convert_local(path, **opts)
    except Exception as exc:  # noqa: BLE001
        friendly = _friendly_office_error(path, exc)
        if friendly:
            raise ConversionError(friendly) from exc
        raise ConversionError(f"转换失败：{exc}") from exc

    if progress:
        progress("正在清理结果…")
    return _finalize_markdown(path, _result_text(result), light=True)


def convert_local_path(
    source: str | Path,
    *,
    progress: Callable[[str], None] | None = None,
    keep_data_uris: bool | None = None,
) -> str:
    """窄接口：仅 convert_local，不经过 Excel/Notebook/ZIP 本地增强。"""
    path = Path(source)
    if not path.exists():
        raise ConversionError(f"文件不存在：{source}")
    if not path.is_file():
        raise ConversionError(f"不是有效文件：{source}")
    enforce_precheck(str(path))
    if progress:
        progress(f"窄接口模式：{path.name}")
    return _run_markitdown_local(path, progress=progress, keep_data_uris=keep_data_uris)


def convert_stream(
    stream: BinaryIO,
    *,
    extension: str | None = None,
    mime_type: str | None = None,
    charset: str | None = None,
    progress: Callable[[str], None] | None = None,
    keep_data_uris: bool | None = None,
) -> str:
    """从二进制流转换（用于 stdin 管道）。"""
    try:
        from markitdown import StreamInfo
    except ImportError as exc:
        raise ConversionError("未安装转换引擎。") from exc

    ext = (extension or "").strip()
    if ext and not ext.startswith("."):
        ext = "." + ext

    stream_info = None
    if ext or mime_type or charset:
        stream_info = StreamInfo(
            extension=ext or None,
            mimetype=(mime_type or None),
            charset=(charset or None),
        )

    if progress:
        progress("正在从标准输入读取…")

    if not stream.seekable():
        buffer = io.BytesIO()
        while True:
            chunk = stream.read(4096)
            if not chunk:
                break
            buffer.write(chunk)
        buffer.seek(0)
        stream = buffer

    md = _get_markitdown()
    opts = _merge_convert_opts(keep_data_uris)
    try:
        if progress:
            progress("正在转换标准输入流…")
        result = md.convert_stream(stream, stream_info=stream_info, **opts)
    except Exception as exc:  # noqa: BLE001
        raise ConversionError(f"转换失败：{exc}") from exc

    if progress:
        progress("正在清理结果…")
    return _finalize_markdown(None, _result_text(result), light=True)


def convert_path(
    source: str | Path,
    *,
    progress: Callable[[str], None] | None = None,
    local_only: bool | None = None,
    keep_data_uris: bool | None = None,
) -> str:
    """将本地文件或 URL 转为 Markdown。"""
    if isinstance(source, Path):
        source = str(source)
    source = (source or "").strip()
    if not source:
        raise ConversionError("请选择文件或输入 URL。")

    narrow = _should_use_local_only(local_only)
    keep = _resolve_keep_data_uris(keep_data_uris)

    if progress:
        progress("正在初始化转换引擎…")

    is_url = source.startswith(("http://", "https://"))
    path: Path | None = None
    if not is_url:
        path = Path(source)
        if not path.exists():
            raise ConversionError(f"文件不存在：{source}")
        if not path.is_file():
            raise ConversionError(f"不是有效文件：{source}")

        enforce_precheck(str(path))

        if narrow:
            return convert_local_path(path, progress=progress, keep_data_uris=keep_data_uris)

        if progress:
            progress(f"正在转换：{path.name}")

        if path.suffix.lower() in {".xlsx", ".xls"}:
            try:
                if progress:
                    progress("正在解析 Excel（合并单元格 / 分区域）…")
                from excel_convert import convert_excel_to_markdown

                text = convert_excel_to_markdown(path)
                text = clean_markdown(text, keep_data_uris=keep)
                if progress:
                    progress("转换完成")
                return text
            except Exception as exc:  # noqa: BLE001
                if progress:
                    progress(f"Excel 增强转换失败，回退标准引擎…（{exc}）")

        if path.suffix.lower() == ".ipynb":
            try:
                if progress:
                    progress("正在解析 Jupyter Notebook（代码输出 / 表格）…")
                from ipynb_convert import convert_ipynb_to_markdown

                text = convert_ipynb_to_markdown(path)
                text = clean_markdown(text, keep_data_uris=keep)
                if progress:
                    progress("转换完成")
                return text
            except Exception as exc:  # noqa: BLE001
                if progress:
                    progress(f"Notebook 增强转换失败，回退标准引擎…（{exc}）")

        if path.suffix.lower() == ".zip":
            try:
                if progress:
                    progress("正在解析 ZIP…")
                from zip_convert import convert_zip_to_markdown

                text = convert_zip_to_markdown(path, progress=progress)
                text = clean_markdown(text, keep_data_uris=keep)
                if progress:
                    progress("转换完成")
                return text
            except ValueError as exc:
                raise ConversionError(str(exc)) from exc
            except Exception as exc:  # noqa: BLE001
                if progress:
                    progress(f"ZIP 增强转换失败，回退标准引擎…（{exc}）")
    else:
        if progress:
            progress(f"正在转换 URL：{source}")

    if progress:
        progress("正在加载转换引擎（首次可能较慢）…")
    md = _get_markitdown()
    opts = _merge_convert_opts(keep_data_uris)
    try:
        if progress:
            progress(f"正在转换：{Path(source).name if path else source}")
        result = md.convert(source, **opts)
    except Exception as exc:  # noqa: BLE001
        msg = str(exc)
        if "Max retries" in msg or "ConnectTimeout" in msg or "Connection" in msg:
            raise ConversionError(
                f"网络请求失败，请检查网络或稍后重试。\n详情：{exc}"
            ) from exc
        if "404" in msg:
            raise ConversionError(f"资源不存在（404）。\n详情：{exc}") from exc
        if path is not None:
            friendly = _friendly_office_error(path, exc)
            if friendly:
                raise ConversionError(friendly) from exc
            if path.suffix.lower() == ".zip":
                raise ConversionError(
                    "ZIP 转换失败。若扩展名是 .zip 但实际是 RAR/7z，请先改为真正的 ZIP。\n"
                    f"详情：{exc}"
                ) from exc
        raise ConversionError(f"转换失败：{exc}") from exc

    text = _result_text(result)

    if progress:
        progress("正在清理结果…")
    text = _finalize_markdown(path, text, light=True)

    if progress:
        progress("转换完成")
    return text


def default_output_path(source: str) -> Path:
    if source.startswith(("http://", "https://")):
        return Path.cwd() / "converted.md"
    return Path(source).with_suffix(".md")


def is_supported_file(path: str | Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS
